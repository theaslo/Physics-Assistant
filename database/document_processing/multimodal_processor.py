#!/usr/bin/env python3
"""
Multimodal Document Processor for Physics Educational Content
Processes PDFs, text documents, and mixed content with equations and diagrams.
"""
import os
import logging
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass
import json
import hashlib
from datetime import datetime
import tempfile
import shutil

# Document processing libraries
import PyPDF2
import docx
from PIL import Image
import fitz  # PyMuPDF for better PDF processing

# Our custom processors
from latex_processor import PhysicsLatexProcessor, LatexEquation
from diagram_analyzer import PhysicsDiagramAnalyzer, PhysicsDiagram

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentSection:
    """Represents a section of a document with multimodal content"""
    section_id: str
    section_type: str  # text, image, equation, mixed
    content: str
    page_number: Optional[int]
    coordinates: Optional[Tuple[int, int, int, int]]  # x, y, width, height
    metadata: Dict[str, Any]

@dataclass
class ExtractedContent:
    """Represents all content extracted from a document"""
    text_content: str
    equations: List[LatexEquation]
    diagrams: List[PhysicsDiagram]
    images: List[str]  # paths to extracted images
    sections: List[DocumentSection]
    metadata: Dict[str, Any]

@dataclass
class ProcessedDocument:
    """Complete processed document with all analysis"""
    file_path: str
    document_hash: str
    content: ExtractedContent
    educational_classification: Dict[str, Any]
    physics_concepts: List[str]
    difficulty_level: str
    processing_timestamp: datetime
    processing_metadata: Dict[str, Any]

class MultimodalDocumentProcessor:
    """Advanced processor for multimodal physics educational documents"""
    
    def __init__(self, output_dir: str = None):
        self.latex_processor = PhysicsLatexProcessor()
        self.diagram_analyzer = PhysicsDiagramAnalyzer()
        self.output_dir = output_dir or tempfile.mkdtemp(prefix="physics_docs_")
        
        # Educational content patterns
        self.problem_patterns = [
            r'problem\s+\d+',
            r'exercise\s+\d+',
            r'question\s+\d+',
            r'find\s+the',
            r'calculate\s+the',
            r'determine\s+the',
            r'what\s+is\s+the',
            r'a\s+\d+\s*kg\s+',  # Mass problems
            r'a\s+car\s+',       # Vehicle problems
            r'a\s+block\s+',     # Block problems
        ]
        
        self.solution_patterns = [
            r'solution:?',
            r'answer:?',
            r'given:?.*find:?',
            r'step\s+\d+',
            r'therefore',
            r'thus',
            r'substituting',
        ]
        
        self.explanation_patterns = [
            r'explanation:?',
            r'concept:?',
            r'note\s+that',
            r'remember\s+that',
            r'this\s+means',
            r'in\s+other\s+words',
        ]
        
        # Physics terminology for concept extraction
        self.physics_terms = {
            'mechanics': ['force', 'velocity', 'acceleration', 'momentum', 'energy', 'work', 'power'],
            'waves': ['frequency', 'wavelength', 'amplitude', 'wave', 'oscillation', 'harmonic'],
            'thermodynamics': ['heat', 'temperature', 'entropy', 'gas', 'pressure', 'volume'],
            'electromagnetism': ['electric', 'magnetic', 'field', 'charge', 'current', 'voltage'],
            'optics': ['light', 'lens', 'mirror', 'reflection', 'refraction', 'ray']
        }
        
        logger.info(f"Multimodal document processor initialized with output dir: {self.output_dir}")
    
    def process_document(self, file_path: str) -> ProcessedDocument:
        """Main method to process any supported document type"""
        try:
            logger.info(f"Processing document: {file_path}")
            
            # Validate file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            # Calculate document hash for caching/deduplication
            doc_hash = self._calculate_file_hash(file_path)
            
            # Extract content based on file type
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                content = self._process_pdf(file_path)
            elif file_ext == '.docx':
                content = self._process_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                content = self._process_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Analyze and classify content
            educational_classification = self._classify_educational_content(content)
            physics_concepts = self._extract_physics_concepts(content)
            difficulty_level = self._assess_difficulty_level(content, educational_classification)
            
            # Create processing metadata
            processing_metadata = {
                'file_size': os.path.getsize(file_path),
                'file_type': file_ext,
                'extraction_stats': {
                    'text_length': len(content.text_content),
                    'equations_count': len(content.equations),
                    'diagrams_count': len(content.diagrams),
                    'images_count': len(content.images),
                    'sections_count': len(content.sections)
                }
            }
            
            return ProcessedDocument(
                file_path=file_path,
                document_hash=doc_hash,
                content=content,
                educational_classification=educational_classification,
                physics_concepts=physics_concepts,
                difficulty_level=difficulty_level,
                processing_timestamp=datetime.now(),
                processing_metadata=processing_metadata
            )
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file for deduplication"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _process_pdf(self, file_path: str) -> ExtractedContent:
        """Process PDF documents with advanced text and image extraction"""
        text_content = ""
        equations = []
        diagrams = []
        images = []
        sections = []
        
        # Use PyMuPDF for better PDF processing
        pdf_document = fitz.open(file_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Extract text
            page_text = page.get_text()
            text_content += page_text + "\n"
            
            # Create text section
            if page_text.strip():
                sections.append(DocumentSection(
                    section_id=f"page_{page_num}_text",
                    section_type="text",
                    content=page_text,
                    page_number=page_num + 1,
                    coordinates=None,
                    metadata={"extraction_method": "pymupdf"}
                ))
            
            # Extract images from page
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_document, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_path = os.path.join(
                            self.output_dir,
                            f"page_{page_num}_img_{img_index}.png"
                        )
                        pix.save(img_path)
                        images.append(img_path)
                        
                        # Analyze image if it looks like a physics diagram
                        if self._is_likely_physics_diagram(img_path):
                            diagram = self.diagram_analyzer.analyze_image(img_path)
                            diagrams.append(diagram)
                            
                            sections.append(DocumentSection(
                                section_id=f"page_{page_num}_diagram_{img_index}",
                                section_type="image",
                                content=f"Physics diagram: {diagram.diagram_type}",
                                page_number=page_num + 1,
                                coordinates=None,
                                metadata={
                                    "image_path": img_path,
                                    "diagram_type": diagram.diagram_type,
                                    "complexity": diagram.complexity_score
                                }
                            ))
                    
                    pix = None
                    
                except Exception as e:
                    logger.warning(f"Error extracting image {img_index} from page {page_num}: {str(e)}")
        
        pdf_document.close()
        
        # Process equations from all text
        equations = self.latex_processor.process_document_equations(text_content)
        
        # Add equation sections
        for i, equation in enumerate(equations):
            sections.append(DocumentSection(
                section_id=f"equation_{i}",
                section_type="equation",
                content=equation.original_latex,
                page_number=None,
                coordinates=None,
                metadata={
                    "domain": equation.physics_domain,
                    "complexity": equation.complexity_score,
                    "variables": equation.variables
                }
            ))
        
        metadata = {
            "total_pages": len(pdf_document),
            "extraction_method": "pymupdf",
            "images_extracted": len(images),
            "diagrams_found": len(diagrams)
        }
        
        return ExtractedContent(
            text_content=text_content,
            equations=equations,
            diagrams=diagrams,
            images=images,
            sections=sections,
            metadata=metadata
        )
    
    def _process_docx(self, file_path: str) -> ExtractedContent:
        """Process Word documents"""
        text_content = ""
        sections = []
        
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        for para_num, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
                
                sections.append(DocumentSection(
                    section_id=f"paragraph_{para_num}",
                    section_type="text",
                    content=paragraph.text,
                    page_number=None,
                    coordinates=None,
                    metadata={"style": paragraph.style.name if paragraph.style else "Normal"}
                ))
        
        # Extract equations
        equations = self.latex_processor.process_document_equations(text_content)
        
        # TODO: Extract embedded images from docx
        images = []
        diagrams = []
        
        metadata = {
            "paragraphs_count": len(doc.paragraphs),
            "extraction_method": "python-docx"
        }
        
        return ExtractedContent(
            text_content=text_content,
            equations=equations,
            diagrams=diagrams,
            images=images,
            sections=sections,
            metadata=metadata
        )
    
    def _process_text(self, file_path: str) -> ExtractedContent:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        # Process equations
        equations = self.latex_processor.process_document_equations(text_content)
        
        # Split into sections (by paragraphs or chapters)
        sections = []
        paragraphs = text_content.split('\n\n')
        
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                sections.append(DocumentSection(
                    section_id=f"paragraph_{i}",
                    section_type="text",
                    content=paragraph,
                    page_number=None,
                    coordinates=None,
                    metadata={}
                ))
        
        metadata = {
            "paragraphs_count": len(paragraphs),
            "extraction_method": "plain_text"
        }
        
        return ExtractedContent(
            text_content=text_content,
            equations=equations,
            diagrams=[],
            images=[],
            sections=sections,
            metadata=metadata
        )
    
    def _is_likely_physics_diagram(self, image_path: str) -> bool:
        """Determine if an image is likely to be a physics diagram"""
        try:
            # Basic heuristics based on image properties
            img = Image.open(image_path)
            width, height = img.size
            
            # Physics diagrams are usually not too small or too large
            if width < 50 or height < 50 or width > 2000 or height > 2000:
                return False
            
            # Check aspect ratio (physics diagrams usually have reasonable ratios)
            aspect_ratio = width / height
            if aspect_ratio < 0.2 or aspect_ratio > 5:
                return False
            
            # TODO: Add more sophisticated checks
            # - Color analysis (diagrams often have limited colors)
            # - Edge detection (diagrams have clear lines)
            # - Text presence (diagrams often have labels)
            
            return True
            
        except Exception:
            return False
    
    def _classify_educational_content(self, content: ExtractedContent) -> Dict[str, Any]:
        """Classify the educational nature of the content"""
        import re
        
        text = content.text_content.lower()
        
        # Count patterns
        problem_count = sum(1 for pattern in self.problem_patterns if re.search(pattern, text))
        solution_count = sum(1 for pattern in self.solution_patterns if re.search(pattern, text))
        explanation_count = sum(1 for pattern in self.explanation_patterns if re.search(pattern, text))
        
        # Classify primary type
        total_indicators = problem_count + solution_count + explanation_count
        
        if total_indicators == 0:
            primary_type = "reference_material"
        elif problem_count > solution_count and problem_count > explanation_count:
            primary_type = "problem_set"
        elif solution_count > problem_count and solution_count > explanation_count:
            primary_type = "solution_manual"
        elif explanation_count > problem_count and explanation_count > solution_count:
            primary_type = "textbook_chapter"
        else:
            primary_type = "mixed_educational"
        
        # Determine if it's worked examples
        has_worked_examples = problem_count > 0 and solution_count > 0
        
        # Check for step-by-step solutions
        step_patterns = [r'step\s+\d+', r'first,?', r'second,?', r'then,?', r'finally,?']
        has_step_by_step = any(re.search(pattern, text) for pattern in step_patterns)
        
        return {
            "primary_type": primary_type,
            "has_problems": problem_count > 0,
            "has_solutions": solution_count > 0,
            "has_explanations": explanation_count > 0,
            "has_worked_examples": has_worked_examples,
            "has_step_by_step": has_step_by_step,
            "pattern_counts": {
                "problems": problem_count,
                "solutions": solution_count,
                "explanations": explanation_count
            },
            "equation_count": len(content.equations),
            "diagram_count": len(content.diagrams)
        }
    
    def _extract_physics_concepts(self, content: ExtractedContent) -> List[str]:
        """Extract physics concepts mentioned in the content"""
        import re
        
        concepts = set()
        text = content.text_content.lower()
        
        # Extract concepts from physics terms
        for domain, terms in self.physics_terms.items():
            for term in terms:
                if re.search(rf'\b{re.escape(term)}\b', text):
                    concepts.add(term)
        
        # Extract concepts from equations
        for equation in content.equations:
            if equation.physics_domain:
                concepts.add(equation.physics_domain)
            # Add concepts based on variables
            for var in equation.variables:
                if var in ['F', 'force']:
                    concepts.add('force')
                elif var in ['v', 'velocity']:
                    concepts.add('velocity')
                elif var in ['a', 'acceleration']:
                    concepts.add('acceleration')
                elif var in ['E', 'energy']:
                    concepts.add('energy')
        
        # Extract concepts from diagrams
        for diagram in content.diagrams:
            concepts.update(diagram.physics_concepts)
        
        return sorted(list(concepts))
    
    def _assess_difficulty_level(self, content: ExtractedContent, classification: Dict[str, Any]) -> str:
        """Assess the difficulty level of the educational content"""
        score = 0
        
        # Base score from content type
        if classification["primary_type"] == "problem_set":
            score += 2
        elif classification["primary_type"] == "solution_manual":
            score += 3
        elif classification["primary_type"] == "textbook_chapter":
            score += 1
        
        # Score from equations
        equation_complexity = sum(eq.complexity_score for eq in content.equations)
        if equation_complexity > 50:
            score += 3
        elif equation_complexity > 20:
            score += 2
        elif equation_complexity > 10:
            score += 1
        
        # Score from diagrams
        diagram_complexity = sum(diag.complexity_score for diag in content.diagrams)
        if diagram_complexity > 30:
            score += 2
        elif diagram_complexity > 15:
            score += 1
        
        # Score from advanced physics terms
        advanced_terms = ['relativistic', 'quantum', 'thermodynamic', 'electromagnetic', 'differential']
        text_lower = content.text_content.lower()
        advanced_count = sum(1 for term in advanced_terms if term in text_lower)
        score += advanced_count
        
        # Classify difficulty
        if score >= 8:
            return "advanced"
        elif score >= 4:
            return "intermediate"
        else:
            return "beginner"
    
    def save_processed_document(self, doc: ProcessedDocument, output_path: str = None) -> str:
        """Save processed document to JSON file"""
        if output_path is None:
            filename = f"processed_{doc.document_hash[:8]}.json"
            output_path = os.path.join(self.output_dir, filename)
        
        # Convert to serializable format
        doc_dict = {
            "file_path": doc.file_path,
            "document_hash": doc.document_hash,
            "processing_timestamp": doc.processing_timestamp.isoformat(),
            "educational_classification": doc.educational_classification,
            "physics_concepts": doc.physics_concepts,
            "difficulty_level": doc.difficulty_level,
            "processing_metadata": doc.processing_metadata,
            "content": {
                "text_length": len(doc.content.text_content),
                "text_preview": doc.content.text_content[:500] + "..." if len(doc.content.text_content) > 500 else doc.content.text_content,
                "equations": [
                    {
                        "original_latex": eq.original_latex,
                        "cleaned_latex": eq.cleaned_latex,
                        "variables": eq.variables,
                        "equation_type": eq.equation_type,
                        "physics_domain": eq.physics_domain,
                        "complexity_score": eq.complexity_score,
                        "is_valid": eq.is_valid
                    }
                    for eq in doc.content.equations
                ],
                "diagrams": [
                    {
                        "image_path": diag.image_path,
                        "diagram_type": diag.diagram_type,
                        "objects": diag.objects,
                        "physics_concepts": diag.physics_concepts,
                        "complexity_score": diag.complexity_score,
                        "vector_count": len(diag.vectors)
                    }
                    for diag in doc.content.diagrams
                ],
                "sections": [
                    {
                        "section_id": sect.section_id,
                        "section_type": sect.section_type,
                        "content_preview": sect.content[:200] + "..." if len(sect.content) > 200 else sect.content,
                        "page_number": sect.page_number,
                        "metadata": sect.metadata
                    }
                    for sect in doc.content.sections
                ],
                "metadata": doc.content.metadata
            }
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(doc_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Processed document saved to: {output_path}")
        return output_path
    
    def get_processing_summary(self, doc: ProcessedDocument) -> Dict[str, Any]:
        """Get a summary of the document processing results"""
        return {
            "file_info": {
                "path": doc.file_path,
                "hash": doc.document_hash[:16] + "...",
                "size": doc.processing_metadata.get("file_size", 0),
                "type": doc.processing_metadata.get("file_type", "unknown")
            },
            "content_summary": {
                "text_length": len(doc.content.text_content),
                "equations_found": len(doc.content.equations),
                "diagrams_found": len(doc.content.diagrams),
                "images_extracted": len(doc.content.images),
                "sections_created": len(doc.content.sections)
            },
            "analysis": {
                "educational_type": doc.educational_classification["primary_type"],
                "difficulty_level": doc.difficulty_level,
                "physics_concepts": doc.physics_concepts[:10],  # Top 10
                "has_worked_examples": doc.educational_classification.get("has_worked_examples", False)
            },
            "processing": {
                "timestamp": doc.processing_timestamp.isoformat(),
                "success": True
            }
        }

# Example usage and testing
if __name__ == "__main__":
    processor = MultimodalDocumentProcessor()
    
    print("Multimodal Document Processor initialized")
    print(f"Output directory: {processor.output_dir}")
    print("Supported formats: PDF, DOCX, TXT, MD")
    print("Features: LaTeX equations, physics diagrams, educational classification")